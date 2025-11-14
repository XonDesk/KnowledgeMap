#!/usr/bin/env python3
"""
Text file handler for file_dispatcher.py
Processes .txt, .md, .markdown, and .doc/.docx files.
Converts them to markdown and saves to library/ folder.
"""
import sys
import subprocess
import logging
import os
from pathlib import Path
import shutil
from datetime import datetime
try:
    from docx import Document
except ImportError:
    import warnings
    warnings.warn(
        "The 'docx' library (python-docx) is not installed or is outdated. "
        "Please install or upgrade it using 'pip install python-docx'. "
        "Some functionalities for .docx files might be limited."
    )
    Document = None # Define Document as None to avoid NameError if not imported

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)


def strip_word_formatting(docx_path):
    """
    Extract plain text from a Word document, removing all formatting.
    
    Args:
        docx_path: Path to the .docx file
        
    Returns:
        Plain text content as string
    """
    logging.info(f"Extracting text from Word document: {docx_path}")
    try:
        from docx import Document
    except ImportError:
        logging.error("python-docx module not installed. Install with: pip install python-docx")
        raise
    
    doc = Document(docx_path)
    text_content = []
    
    paragraph_count = 0
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
        paragraph_count += 1
    
    # Also extract text from tables
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                text_content.append(cell.text)
    
    logging.info(f"Extracted {paragraph_count} paragraphs and {table_count} tables")
    return '\n'.join(text_content)


def process_file(file_path):
    """
    Process a single file: convert to markdown and save to library/ folder.
    
    Args:
        file_path: Path to the input file
        
    Returns:
        Path to the output file
    """
    file_path = Path(file_path)
    logging.info(f"Starting to process file: {file_path}")
    
    if not file_path.exists():
        logging.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Log file details
    file_size = file_path.stat().st_size
    logging.info(f"File size: {file_size:,} bytes")
    
    # Create library directory in the same folder as the script
    script_dir = Path(__file__).parent
    library_path = script_dir / 'library'
    if not library_path.exists():
        logging.info(f"Creating library directory: {library_path}")
        library_path.mkdir(exist_ok=True)
    
    # Determine output filename
    output_filename = file_path.stem + '.md'
    output_path = library_path / output_filename
    logging.info(f"Output file will be: {output_path}")
    
    # Check if output file already exists
    if output_path.exists():
        logging.info(f"Output file already exists, will be overwritten: {output_path}")
    
    # Process based on file type
    file_extension = file_path.suffix.lower()
    logging.info(f"Detected file extension: {file_extension}")
    
    if file_extension in ['.md', '.markdown']:
        # Copy markdown files as-is
        logging.info(f"File is already markdown format, copying as-is")
        shutil.copy2(file_path, output_path)
        logging.info(f"Successfully copied markdown: {file_path.name} -> library/{output_filename}")
        print(f"Copied markdown: {file_path.name} -> library/{output_filename}")
        
    elif file_extension in ['.docx', '.doc']:
        # Extract text from Word document and save as markdown
        logging.info(f"Processing Word document: {file_extension}")
        try:
            text_content = strip_word_formatting(file_path)
            logging.info(f"Extracted {len(text_content)} characters from Word document")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            logging.info(f"Successfully converted Word document: {file_path.name} -> library/{output_filename}")
            print(f"Converted Word document: {file_path.name} -> library/{output_filename}")
        except ImportError as e:
            logging.error(f"Cannot process Word document - missing dependency: {e}")
            logging.error(f"SKIPPED: {file_path.name} - Reason: python-docx module not installed")
            raise
        
    elif file_extension == '.txt':
        # Read text file and save as markdown
        logging.info(f"Processing plain text file")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            logging.info(f"Read {len(text_content)} characters from text file")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            logging.info(f"Successfully converted text file: {file_path.name} -> library/{output_filename}")
            print(f"Converted text file: {file_path.name} -> library/{output_filename}")
        except UnicodeDecodeError as e:
            logging.error(f"Cannot read file - encoding error: {e}")
            logging.error(f"SKIPPED: {file_path.name} - Reason: File encoding not supported (not UTF-8)")
            raise
        
    else:
        logging.error(f"Unsupported file type: {file_extension}")
        logging.error(f"SKIPPED: {file_path.name} - Reason: File extension '{file_extension}' not supported (only .txt, .md, .markdown, .doc, .docx)")
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    logging.info(f"File processing complete: {file_path.name}")
    return output_path


def main():
    logging.info("=" * 60)
    logging.info("Text Handler Started")
    logging.info("=" * 60)
    
    if len(sys.argv) < 2:
        logging.error("No file path provided as argument")
        print("Error: No file path provided", file=sys.stderr)
        sys.exit(1)
    
    file_path = sys.argv[1]
    logging.info(f"Received file path argument: {file_path}")
    
    try:
        output_path = process_file(file_path)
        
        logging.info("=" * 60)
        logging.info("Text Handler Completed Successfully")
        logging.info("=" * 60)
        sys.exit(0)
    except Exception as e:
        logging.error("=" * 60)
        logging.error(f"Text Handler Failed: {e}")
        logging.error("=" * 60)
        logging.exception("Full traceback:")
        print(f"Error processing file: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
